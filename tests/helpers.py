from __future__ import annotations

import csv
import shutil
from contextlib import contextmanager
from pathlib import Path
from uuid import uuid4

from docx import Document
from openpyxl import Workbook


LEGAL_SAMPLE_LINES = (
    "LEI COMPLEMENTAR Nº 999, DE 1º DE JANEIRO DE 2026",
    "O PRESIDENTE DA REPÚBLICA Faço saber que o Congresso Nacional decreta e eu sanciono a seguinte Lei Complementar:",
    "PARTE GERAL",
    "DAS DISPOSIÇÕES PRELIMINARES",
    "LIVRO I",
    "DISPOSIÇÕES GERAIS",
    "TÍTULO I",
    "DAS NORMAS INICIAIS",
    "CAPÍTULO I",
    "DA ORGANIZAÇÃO",
    "Seção I",
    "Das Regras Básicas",
    "Subseção I",
    "Da Estrutura Inicial",
    "Art. 1º Esta lei estabelece normas gerais.",
    "§ 1º O disposto neste artigo aplica-se a todos.",
    "I - observar a legalidade;",
    "a) cumprir requisitos mínimos;",
    "1. registrar atos essenciais;",
    "Art. 2º Ficam revogadas as disposições em contrário.",
)
AMENDING_SAMPLE_LINES = (
    "EMENDA CONSTITUCIONAL Nº 132, DE 20 DE DEZEMBRO DE 2023",
    "As Mesas da Câmara dos Deputados e do Senado Federal promulgam a seguinte Emenda ao texto constitucional:",
    "Art. 1º Os arts. 43, 50 e 105 da Constituição Federal passam a vigorar com as seguintes alterações:",
    "Art. 43. Compete à lei complementar disciplinar aspectos gerais do sistema.",
    "§ 4º Lei complementar poderá estabelecer normas específicas de coordenação.",
    "Art. 50. O Congresso Nacional e suas Casas terão competência para fiscalizar a execução.",
    "Art. 105. Compete ao Superior Tribunal de Justiça:",
    "I - processar e julgar, originariamente, os conflitos de competência;",
    "j) conflitos entre autoridades administrativas e judiciais relacionados ao novo regime;",
    "Art. 2º Esta Emenda Constitucional entra em vigor na data de sua publicação.",
)
MANUAL_DOCX_BLOCKS = (
    ("Heading 1", "MANUAL OPERACIONAL DE APURAÇÃO"),
    ("Heading 2", "Primeiros Passos"),
    ("Heading 3", "Objetivos"),
    (None, "Apresentar o fluxo inicial de conferência do arquivo."),
    ("Heading 3", "Procedimento"),
    ("Heading 4", "Etapa 1 - Receber arquivo"),
    (None, "Verificar extensão, assinatura e integridade."),
    ("Heading 4", "Etapa 2 - Validar conteúdo"),
    (None, "Conferir campos obrigatórios e mensagens de erro."),
    ("Heading 2", "Resumo"),
    (None, "Registrar o resultado final da análise no sistema."),
)


def create_legal_docx(path: Path) -> Path:
    document = Document()
    for paragraph_text in LEGAL_SAMPLE_LINES:
        document.add_paragraph(paragraph_text)

    document.save(path)
    return path


def create_legal_txt(path: Path) -> Path:
    path.write_text("\n".join(LEGAL_SAMPLE_LINES), encoding="utf-8")
    return path


def create_legal_html(path: Path) -> Path:
    lines = [
        "<!DOCTYPE html>",
        "<html lang=\"pt-BR\">",
        "<head>",
        "<meta charset=\"utf-8\">",
        "<title>Lei Complementar nº 999</title>",
        "</head>",
        "<body>",
    ]
    for line in LEGAL_SAMPLE_LINES:
        lines.append(f"<p>{line}</p>")
    lines.extend(["</body>", "</html>"])
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def create_legal_xml(path: Path) -> Path:
    lines = ["<documento>"]
    for index, line in enumerate(LEGAL_SAMPLE_LINES, start=1):
        lines.append(f"  <bloco ordem=\"{index}\">{line}</bloco>")
    lines.append("</documento>")
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def create_legal_csv(path: Path) -> Path:
    with path.open("w", encoding="utf-8", newline="") as file_obj:
        writer = csv.writer(file_obj)
        for line in LEGAL_SAMPLE_LINES:
            writer.writerow([line])
    return path


def create_legal_xlsx(path: Path) -> Path:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Norma"
    for row_number, line in enumerate(LEGAL_SAMPLE_LINES, start=1):
        sheet.cell(row=row_number, column=1).value = line
    workbook.save(path)
    return path


def create_amending_docx(path: Path) -> Path:
    document = Document()
    for paragraph_text in AMENDING_SAMPLE_LINES:
        document.add_paragraph(paragraph_text)

    document.save(path)
    return path


def create_amending_txt(path: Path) -> Path:
    path.write_text("\n".join(AMENDING_SAMPLE_LINES), encoding="utf-8")
    return path


def create_amending_html(path: Path) -> Path:
    lines = [
        "<!DOCTYPE html>",
        "<html lang=\"pt-BR\">",
        "<head>",
        "<meta charset=\"utf-8\">",
        "<title>Emenda Constitucional nº 132</title>",
        "</head>",
        "<body>",
    ]
    for line in AMENDING_SAMPLE_LINES:
        lines.append(f"<p>{line}</p>")
    lines.extend(["</body>", "</html>"])
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def create_amending_csv(path: Path) -> Path:
    with path.open("w", encoding="utf-8", newline="") as file_obj:
        writer = csv.writer(file_obj)
        for line in AMENDING_SAMPLE_LINES:
            writer.writerow([line])
    return path


def create_amending_xlsx(path: Path) -> Path:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Emenda"
    for row_number, line in enumerate(AMENDING_SAMPLE_LINES, start=1):
        sheet.cell(row=row_number, column=1).value = line
    workbook.save(path)
    return path


def create_manual_docx(path: Path) -> Path:
    document = Document()
    for style_name, paragraph_text in MANUAL_DOCX_BLOCKS:
        paragraph = document.add_paragraph(paragraph_text)
        if style_name:
            paragraph.style = style_name

    document.save(path)
    return path


@contextmanager
def workspace_test_dir() -> Path:
    root = Path.cwd() / "_test_artifacts" / uuid4().hex
    root.mkdir(parents=True, exist_ok=True)
    try:
        yield root
    finally:
        shutil.rmtree(root, ignore_errors=True)
